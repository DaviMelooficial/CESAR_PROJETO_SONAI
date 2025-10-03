"""
Script para extração de texto e dados de arquivos DOC/DOCX
Autor: Sistema de Análise e Visualização
Data: Setembro 2025
"""

import os
import json
import csv
from typing import List, Dict, Optional
from pathlib import Path
import logging
from docx import Document
import zipfile
import xml.etree.ElementTree as ET

# Configuração do logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class ExtratorDOC:
    """
    Classe para extrair texto e dados de arquivos DOC/DOCX
    """
    
    def __init__(self, pasta_origem: str = "../data/raw", pasta_destino: str = "../data/processed"):
        self.pasta_origem = Path(pasta_origem)
        self.pasta_destino = Path(pasta_destino)
        
        # Criar pasta de destino se não existir
        self.pasta_destino.mkdir(parents=True, exist_ok=True)
        
    def extrair_docx_python_docx(self, caminho_docx: str) -> Dict:
        """
        Extrai texto e dados usando python-docx
        """
        resultado = {
            'texto': "",
            'paragrafos': [],
            'tabelas': [],
            'metadados': {}
        }
        
        try:
            doc = Document(caminho_docx)
            
            # Metadados do documento
            core_props = doc.core_properties
            resultado['metadados'] = {
                'autor': core_props.author or "",
                'titulo': core_props.title or "",
                'assunto': core_props.subject or "",
                'criado': str(core_props.created) if core_props.created else "",
                'modificado': str(core_props.modified) if core_props.modified else "",
                'num_paragrafos': len(doc.paragraphs),
                'num_tabelas': len(doc.tables)
            }
            
            # Extrair parágrafos
            texto_completo = ""
            for i, paragrafo in enumerate(doc.paragraphs):
                if paragrafo.text.strip():
                    resultado['paragrafos'].append({
                        'numero': i+1,
                        'texto': paragrafo.text.strip(),
                        'estilo': paragrafo.style.name if paragrafo.style else "Normal"
                    })
                    texto_completo += paragrafo.text + "\n"
            
            resultado['texto'] = texto_completo
            
            # Extrair tabelas
            for i, tabela in enumerate(doc.tables):
                dados_tabela = []
                for linha in tabela.rows:
                    linha_dados = []
                    for celula in linha.cells:
                        linha_dados.append(celula.text.strip())
                    dados_tabela.append(linha_dados)
                
                if dados_tabela:  # Só adicionar se houver dados
                    resultado['tabelas'].append({
                        'tabela_num': i+1,
                        'linhas': len(dados_tabela),
                        'colunas': len(dados_tabela[0]) if dados_tabela else 0,
                        'dados': dados_tabela
                    })
                    
        except Exception as e:
            logger.error(f"Erro ao extrair dados com python-docx de {caminho_docx}: {e}")
            
        return resultado
    
    def extrair_docx_zipfile(self, caminho_docx: str) -> str:
        """
        Extrai texto usando zipfile (fallback method)
        """
        try:
            with zipfile.ZipFile(caminho_docx, 'r') as docx:
                # Ler o document.xml
                xml_content = docx.read('word/document.xml')
                root = ET.fromstring(xml_content)
                
                # Namespace para DOCX
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }
                
                # Extrair todo o texto
                paragrafos = root.findall('.//w:p', namespaces)
                texto_completo = ""
                
                for para in paragrafos:
                    texto_para = ""
                    for t in para.findall('.//w:t', namespaces):
                        if t.text:
                            texto_para += t.text
                    if texto_para.strip():
                        texto_completo += texto_para + "\n"
                        
                return texto_completo
                
        except Exception as e:
            logger.error(f"Erro ao extrair texto com zipfile de {caminho_docx}: {e}")
            return ""
    
    def processar_arquivo(self, caminho_doc: str, formato_saida: str = "json") -> bool:
        """
        Processa um arquivo DOC/DOCX específico
        
        Args:
            caminho_doc: Caminho para o arquivo DOC/DOCX
            formato_saida: 'txt', 'json' ou 'csv'
        """
        if not os.path.exists(caminho_doc):
            logger.error(f"Arquivo não encontrado: {caminho_doc}")
            return False
            
        nome_arquivo = Path(caminho_doc).stem
        logger.info(f"Processando: {nome_arquivo}")
        
        # Determinar se é DOC ou DOCX
        extensao = Path(caminho_doc).suffix.lower()
        
        if extensao == '.docx':
            dados = self.extrair_docx_python_docx(caminho_doc)
            
            if not dados['texto'].strip():
                # Fallback para método zipfile
                logger.warning("python-docx não conseguiu extrair texto, tentando zipfile...")
                texto_fallback = self.extrair_docx_zipfile(caminho_doc)
                dados['texto'] = texto_fallback
                
        else:
            logger.warning(f"Formato DOC não suportado diretamente: {extensao}")
            logger.info("Para arquivos .doc, converta para .docx primeiro")
            return False
        
        # Salvar conforme formato solicitado
        if formato_saida == "txt":
            arquivo_saida = self.pasta_destino / f"{nome_arquivo}_extraido.txt"
            with open(arquivo_saida, 'w', encoding='utf-8') as f:
                f.write(dados['texto'])
                
        elif formato_saida == "json":
            arquivo_saida = self.pasta_destino / f"{nome_arquivo}_dados.json"
            with open(arquivo_saida, 'w', encoding='utf-8') as f:
                json.dump(dados, f, ensure_ascii=False, indent=2)
                
        elif formato_saida == "csv":
            # Salvar parágrafos como CSV
            arquivo_saida = self.pasta_destino / f"{nome_arquivo}_paragrafos.csv"
            
            with open(arquivo_saida, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(['numero', 'estilo', 'texto'])
                for para in dados['paragrafos']:
                    writer.writerow([para['numero'], para['estilo'], para['texto']])
            
            # Salvar tabelas separadamente se existirem
            if dados['tabelas']:
                for tab_info in dados['tabelas']:
                    arquivo_tabela = self.pasta_destino / f"{nome_arquivo}_tabela_{tab_info['tabela_num']}.csv"
                    with open(arquivo_tabela, 'w', newline='', encoding='utf-8') as f:
                        writer = csv.writer(f)
                        writer.writerows(tab_info['dados'])
        
        logger.info(f"Dados salvos em: {arquivo_saida}")
        return True
    
    def processar_pasta(self, formato_saida: str = "json") -> List[str]:
        """
        Processa todos os documentos DOC/DOCX na pasta de origem
        """
        docs_processados = []
        
        if not self.pasta_origem.exists():
            logger.error(f"Pasta de origem não encontrada: {self.pasta_origem}")
            return docs_processados
            
        # Buscar todos os documentos
        arquivos_doc = list(self.pasta_origem.glob("*.docx")) + list(self.pasta_origem.glob("*.doc"))
        
        if not arquivos_doc:
            logger.warning(f"Nenhum arquivo DOC/DOCX encontrado em: {self.pasta_origem}")
            return docs_processados
            
        logger.info(f"Encontrados {len(arquivos_doc)} arquivos DOC/DOCX")
        
        for doc_path in arquivos_doc:
            try:
                if self.processar_arquivo(str(doc_path), formato_saida):
                    docs_processados.append(str(doc_path))
            except Exception as e:
                logger.error(f"Erro ao processar {doc_path}: {e}")
                
        return docs_processados

def main():
    """
    Função principal para execução do script
    """
    print("=== Extrator de Dados DOC/DOCX ===")
    print("1. Processar arquivo específico")
    print("2. Processar todos os DOC/DOCX da pasta")
    
    opcao = input("\nEscolha uma opção (1-2): ").strip()
    
    extrator = ExtratorDOC()
    
    if opcao == "1":
        caminho = input("Digite o caminho do arquivo DOC/DOCX: ").strip()
        formato = input("Formato de saída (txt/json/csv) [json]: ").strip() or "json"
        
        if extrator.processar_arquivo(caminho, formato):
            print("✅ Processamento concluído com sucesso!")
        else:
            print("❌ Erro no processamento.")
            
    elif opcao == "2":
        formato = input("Formato de saída (txt/json/csv) [json]: ").strip() or "json"
        
        processados = extrator.processar_pasta(formato)
        print(f"\n✅ {len(processados)} arquivos processados com sucesso!")
        
        for arquivo in processados:
            print(f"  - {os.path.basename(arquivo)}")
    else:
        print("Opção inválida!")

if __name__ == "__main__":
    main()